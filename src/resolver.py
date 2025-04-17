from pickletools import read_unicodestring1
from tracemalloc import start
from typing import List
from docx import Document
from ollama import Client
import ollama
from pymongo import MongoClient
import xml.etree.ElementTree as ET
import re
import time
from concurrent.futures import ThreadPoolExecutor
from dotenv import load_dotenv
import os
import shutil
from tqdm import tqdm

from utils.llm import call_deepseek,call_ollama
from utils.logger import log_to_mongodb

file_path ="C:\\Lee\\files\\采购\\安阳钢铁集团有限责任公司综利公司烧结机头灰资源化处置项目（运营）\\05 压滤机滤布采购合同.docx"

file_path ="C:\\Lee\\files\\采购\\others\\03 回转窑采购合同.docx"
file_path ="C:\\Lee\\files\\采购\\others\\05 起重设备采购合同.docx"
file_path ="C:\\Lee\\files\\采购\\others\\12低压柜及三箱合同.docx"
file_path ="C:\\Lee\\files\\采购\\安阳钢铁集团有限责任公司综利公司烧结机头灰资源化处置项目（运营）\\04 渣浆泵备件采购合同.docx"
file_path ="C:\\Lee\\files\\03 循环风机采购合同.docx"
file_path ="C:\\Lee\\files\\采购\\others\\04 管道增压泵采购合同.docx"
file_path ="C:\\Lee\\files\\采购\\others\\14自动汽水取样装置采购合同.docx"


table_header_search_system_prompt = '''
你是一位表格数据的识别判断专家，可以帮助分析提供的表格是否为**设备\产品**表格，用户提供表格内容位于<table></table>标签中，识别要求如下：
   1、首先判断是否为**设备\产品**表格：
      设备的表头一般出现在表格的前五行以内；
      每一行数据位于`[]`中；
      常见的设备表格的表头字段通常包含但是不局限有：名称或者设备名称或产品，规格或者规格/材质或者型号，单位，数量，单价，总价，生产厂家等，
      如果主要字段基本符合设备表格的要求，名称或者设备名称是必需的，价格、单价、总价至少有一条是必需的，注意表头字段都是在同一行的。
      如果满足这些字段必须的条件，那么可以判断**是**设备表格,按以下分析策略:
    1.1、需要逐个分析原始表头字段，标准化表头到指定的字段：设备名称、规格/材质、单位、数量、生产厂家、单价、总价；
      （分析的时候注意：一、没有出现价格字段，那么默认单价就是价格；二、标准化字段没有出现在原始表格中，将其结果设置为“ ”空然后原始表格对应提取字段列索引设置为-1。)
    1.2、分析设备表格表头字段生成<head></head>标签：
      分析的时候请遵循以下步骤：
      1.2.1、逐个分析表头字段：
        需要注意如果在分析价格字段的时候，将价格单位进行识别，常规是**元**或者**万元**（如果不属于二者其一，使用表格提供的原始单位数据），价格单位通常出现在价格字段之后，位于（）括号符号内；
      1.2.2、分别整理每个字段生成<head></head>标签；
        表头在表格的行数，用start_row属性表示，注意start_row从0开始计数，0表示第一行，1表示第二行，以此类推。
    将表头的原始字段和索引输出成如下格式生成在<head></head>标签中：
    举例：
    <head>
    [{"字段":"序号";"索引":"0"},{"字段":"名称";"索引":"1"},{"字段":"规格/材质";"索引":"2"},{"字段":"单位";"索引":"3"},{"字段":"数量";"索引":"4"},{"字段":"生产厂家";"索引":"5"},{"字段":"单价";"索引":"6"},{"字段":"总价";"索引":"7"}]
    </head>
    或者
    <head>
    [{"字段":"名称";"索引":"0"},{"字段":"规格/材质";"索引":"1"},{"字段":"单位";"索引":"2"},{"字段":"数量";"索引":"3"},{"字段":"生产厂家";"索引":"4"},{"字段":"单价";"索引":"5"},{"字段":"总价";"索引":"6"},{"字段":"价格单位";"索引":"-1"}]
    </head>
    1.3、 按照<head></head>标签生成XML结构:
      1.3.1、检查是否所有标准字段都被覆盖：
        设备名称：有|无
        规格/材质：有|无
        单位：有|无
        数量：有|无
        生产厂家：有|无
        单价：有|无
        总价：有|无       
        
        最后根据<head></head>标签的内容生成如下格式的XML(start_row属性表示表头的行位置，index属性表示原始表格列索引字段，0表示表头的第一个字段，1表示第二个字段，以此类推)，其中unit属性只有标准字段单价和总价才有:
        <fields start_row="0">
            <field original="名称" standard="设备名称" index="0"/>
            <field original="规格/材质" standard="规格/材质" index="1"/>
            <field original="单位" standard="单位" index="2"/>
            <field original="数量" standard="数量" index="3"/>
            <field original="生产厂家" standard="生产厂家" index="4"/>
            <field original="单价" standard="单价" unit="元" index="5"/>
            <field original="总价" standard="总价" unit="元" index="6"/>
        </fields>
  
   2、如果用户提供的表格**不是**一份设备表格，请直接返回空xml，格式为<xml></xml>，不需要做任何的标准化字段提取。
   
   **注意：请严格遵守以上约定，先识别是否是设备表格，再根据对应的结果进行输出,最终的显示结果仅包含XML。**
   
   在回答中请不要假设任何不属于用户提供的数据，并且不要虚拟数据，如实的根据用户提供的数据回答。
   请step by step的分析用户提供的文本，将分析内容写在<think></think>标签中，\
   再将表头的整理写在<head></head>标签中，\
   最后再综合<think></think>和<head></head>标签中的field的字段的名称对应是否正确，是否与标准化字段的名称匹配，再次检查索引是否是否匹配表头字段。
   再将分析结果输出到独立的<fields></fields>标签中。
   
   以下是用户提供的表格数据：
'''

# 创建键的映射表（中英文对照）
key_mapping = {
    '设备名称': 'device_name',
    '规格/材质': 'specification_material',
    '数量': 'quantity',
    '单位': 'unit',
    '生产厂家': 'manufacturer',
    '单价': 'unit_price',
    '总价': 'total_price',
    '价格单位': 'price_unit',
    'other_column': 'additional_info',
}

ollama_config = {
    'model_name': 'qwq:latest',
    'ollama_host':'http://192.168.43.41:11434',
}

def search_table_header_with_llm(table_content:str):
    
    response = call_deepseek(base_url=deepseek_base_url,api_key=deepseek_api_key,
                             prompt=f"<table>{table_content}</table>",system_prompt=table_header_search_system_prompt,model_name=model_name) 
    return response
    
    # 原始调用
    # client = Client(
    #     host='http://192.168.43.41:11434',
    #     headers={'x-some-header': 'some-value'}
    # )
    # response = client.chat(model=model_name, 
    #                     options={
    #                         'temperature':0,
    #                          "num_ctx": 4096,
    #                     },
    #                     messages=[
    # {
    #     'role': 'user',
    #     'content': f"{table_header_search_system_prompt}<table>{table_content}</table>"+"<think>\n",
    # },
    # ])
    # return response.message.content

def is_valid_data(data_list):
    """
    检查列表中的每个字典是否都包含 'original', 'standard', 'index' 三个键。
    
    参数:
        data_list (list): 嵌套字典的列表
        
    返回:
        bool: 如果所有字典都包含所需键返回 True，否则返回 False
    """
    required_keys = {'original', 'standard','unit', 'index'}
    
    for item in data_list:
        if set(item.keys()) != required_keys:
            return False
    return True

def extract_xml_to_dict(text):
    """
    从文本中提取XML结构并解析为字典列表，同时提取根节点的start_row属性。
    适应新的XML结构，其中field元素通过属性original、standard、index定义。
    
    参数:
        text (str): 包含XML结构的文本
        
    返回:
        tuple: (字段字典列表, start_row值) 如果解析失败返回 (None, None)
    """
    
    if not text:
        return None, None  # 如果文本为空，返回None, None
    
    # 使用正则表达式提取<fields>...</fields>部分
    pattern = r'<fields.*?>.*?</fields>'
    matches = re.findall(pattern, text, re.DOTALL)
    
    if not matches:
        return None, None  # 如果没有找到XML结构，返回None, None
    
    xml_string = matches[-1]  # 取匹配列表中的最后一个元素
    
    # 解析XML
    try:
        root = ET.fromstring(xml_string)
        
        # 提取start_row属性，默认为None如果不存在
        start_row = root.get('start_row', None)
        
        # 解析field元素
        result = []
        for field in root.findall('field'):
            # 直接从属性中提取信息
            field_dict = {
                'original': field.get('original', ''),
                'standard': field.get('standard', ''),
                'unit': field.get('unit', ''),
                'index': field.get('index', '')
            }
            result.append(field_dict)
        
        return result, start_row
    
    except ET.ParseError:
        return None, None  # 如果XML格式错误，返回None, None


def transform_dict(dic):
    # 键的映射关系
    key_mapping = {
        '项目名称': 'project_name',
        '合同编号': 'contract_number',
        '合同类型': 'contract_type',
        '子项名称': 'subitem_name'
    }
    translated_dict = {key_mapping[old_key]: value for old_key, value in dic.items()}
    return translated_dict

def parse_doc_meta_xml_to_dict(text):
    # 使用正则表达式提取最后一个<fields>...</fields>部分
    pattern = r'<fields.*?>.*?</fields>'
    matches = re.findall(pattern, text, re.DOTALL)
    
    if not matches:
        return {"error": "No XML structure found in the text"}
    
    xml_string = matches[-1]  # 取匹配列表中的最后一个元素
    
    # 创建 XML 解析器
    try:
        # 将字符串转换为文件对象并解析
        root = ET.fromstring(xml_string)
        
        # 定义只接受的字段
        allowed_fields = {
            '项目名称': None,
            '合同编号': None,
            '合同类型': None,
            '子项名称': None
        }
        
        # 遍历所有 field 元素
        for field in root.findall('field'):
            attributes = field.attrib
            if len(attributes) == 1:
                key = list(attributes.keys())[0]
                # 只处理指定的四个字段
                if key in allowed_fields:
                    allowed_fields[key] = attributes[key]
                
        return allowed_fields
    
    except ET.ParseError as e:
        return {"error": f"XML parsing error: {str(e)}"}

def parse_table_to_objects(table, mapping_list:List, start_row=0):
    """
    根据映射列表从 python-docx 的 Table 对象中解析数据并转换为对象列表。
    未定义的表头字段将被聚合到 'other_column' 中。
    
    参数:
        table (docx.table.Table): python-docx 的表格对象
        mapping_list (list): 包含字典的列表，每个字典有 'original', 'standard', 'index' 键
        start_row (int): 开始解析的行号，默认为0
        
    返回:
        list: 解析后的对象列表，每个对象是一个字典
    """
    try:
        # 创建从原始标题到标准字段的映射
        header_mapping = {item['original']: item['standard'] for item in mapping_list}
        index_mapping = {item['original']: int(item['index']) for item in mapping_list}
        
        # 获取表头行（第start_row行）
        header_row = table.rows[int(start_row)].cells
        headers = [cell.text.strip() for cell in header_row]
        
        # 验证表头是否与 mapping_list 中的 original 匹配，并找出未定义的表头
        undefined_headers = {}
        for idx, header in enumerate(headers):
            if header not in header_mapping and header != '':
                # print(f"警告: 表头 '{header}' 未在映射列表中定义，将被聚合到 other_column")
                undefined_headers[header] = idx
        
        # 解析数据行
        result = []
        start_row = int(start_row) + 1
        table_row_index = start_row  # 索引行
        for row in table.rows[start_row:]:  # 从start_row+1行开始
            cells = [cell.text.strip() for cell in row.cells]
                
            # 创建对象
            obj = {'table_row_index':table_row_index}
            table_row_index += 1
            
            # 处理已定义的映射字段
            for original, standard in header_mapping.items():
                idx = index_mapping.get(original, -1)
                if idx >= 0 and idx < len(cells):  # 确保索引有效
                    obj[standard] = cells[idx]
            
            # 价格单位字段直接添加
            price_dic_list =  list(filter(lambda x: x['standard']=='单价', mapping_list))
            if len(price_dic_list) > 0:
                obj['价格单位'] = price_dic_list[0]['unit']
            else:
                obj['价格单位'] = ''
            
            # 处理未定义的字段，聚合到 other_column
            if undefined_headers:
                other_column = {}
                for header, idx in undefined_headers.items():
                    if idx < len(cells):  # 确保索引有效
                        other_column[header] = cells[idx]
                if other_column:  # 只有当有内容时才添加 other_column
                    obj['other_column'] = other_column
                    
            
            result.append(obj)
        
        return result
    
    except Exception as e:
        print(f"解析表格数据失败: {str(e)}")
        return []

# 解析文档元数据
def resolve_doc_info(file_path):
    doc = Document(file_path)
    range = 100
    index = 0
    paragraphs = []
    for paragraph in doc.paragraphs:
        if index > range:
            break
        index += 1
        if len(paragraph.text.strip()) > 0:
            paragraphs.append(paragraph.text)
    context = '\n'.join(paragraphs)
    system_prompt='''
        你需要从用户提供的文本中提取出与合同相关的内容，需要提取的字段有：[项目名称, 合同编号, 合同类型, 子项名称]。
        合同编号如果找到多个编号，优先买方合同编号。
        将结果输出为如下xml结构：
        <fields>
            <field 项目名称="名称" />
            <field 合同编号="例如：BGDSYR-2020-12" />
            <field 合同类型="例如：订货|供货|施工|等 合同" />
            <field 子项名称="例如：配件|设备|服务 等" />
        </fields>
        注意不要虚构内容，仅根据用户提供的内容进行数据提取，如果没有找到相关的合同信息，对应的属性值可以为空。
        请step by step的分析用户提供的文本，将分析内容写在<think></think>标签中，再将分析结果输出到独立的<fields></fields>标签中。
    '''
    
    response = call_deepseek(base_url=deepseek_base_url,api_key=deepseek_api_key,
                             prompt=f"{context}",system_prompt=system_prompt,model_name=model_name) 
    return response
    
    # client = Client(
    # host='http://192.168.43.41:11434',
    # headers={'x-some-header': 'some-value'}
    # )
    # response = client.chat(model='qwq:latest', 
    #                     options={
    #                         'temperature':0,
    #                         "num_ctx": 1024*16,
    #                     },
    #                     messages=[
    #     {'role':'system', 'content': system_prompt},
    #     {
    #         'role': 'user',
    #         'content': f"{context}"+"<think>\n",
    #     },
    # ])
    # return response.message.content

# 函数：根据映射表转换字典键
def transform_keys(data, mapping):
    if not isinstance(data, dict):
        return data
    
    transformed = {}
    other_data = {}
    other_column_key = mapping.get('other_column', 'other_column')  # 获取 'other_column' 对应的新键名
    
    for k, v in data.items():
        if k in mapping:  # 如果键在映射表中，使用新键名
            new_key = mapping[k]
            transformed[new_key] = transform_keys(v, mapping)
        else:  # 否则，存入 other_data 待处理
            other_data[k] = transform_keys(v, mapping)
    
    # 如果有未映射的键，将它们存入 other_column 对应的键下
    if other_data:
        if other_column_key in transformed:  # 如果 other_column 已存在，合并数据
            if isinstance(transformed[other_column_key], dict):
                transformed[other_column_key].update(other_data)
            else:  # 如果不是字典（比如是字符串或 None），则覆盖
                transformed[other_column_key] = other_data
        else:  # 如果 other_column 不存在，直接赋值
            transformed[other_column_key] = other_data
    
    return transformed

# 连接到MongoDB
def save_to_mongodb(transformed_data):
    try:
        # 连接到本地MongoDB (根据你的实际配置修改连接字符串)
        client = MongoClient('mongodb://localhost:27017/')
        
        # 选择数据库和集合
        db = client['equipment_db']  # 数据库名
        collection = db['equipment_collection']  # 集合名
        
        # 插入数据
        result = collection.insert_one(transformed_data)
        # print(f"数据插入成功，ID: {result.inserted_id}")
        
        # 关闭连接
        client.close()
    except Exception as e:
        print(f"发生错误: {str(e)}")

def resolve_doc_meta_info_with_llm(file_path:str):
    # 解析文档元数据
    doc_meta_text = resolve_doc_info(file_path)
    doc_meta_dic = parse_doc_meta_xml_to_dict(doc_meta_text)
    doc_meta_dic = transform_dict(doc_meta_dic)
    doc_meta_dic['file_path'] = file_path
    return doc_meta_dic

def retrieve_table_from_docx(file_path:str):
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"打开文件失败: {str(e)}")
        return None
    doc_meta_dic = resolve_doc_meta_info_with_llm(file_path)
    if is_debug:
        print(f"文档元数据\n{doc_meta_dic}")
    
    table_index = 0
    table_dic_list = []
     # 遍历文档中的所有表格
    for table in doc.tables:
        if is_debug:
            print("找到一个表格：\n")
        table_index += 1
        table_context_list = []
        row_index = 0
        # 遍历表格的每一行
        for row in table.rows:
            row_data = []
            # 遍历每一行的每个单元格
            for cell in row.cells:
                row_data.append(cell.text.strip())  # 获取单元格文本并去除多余空格
            # print(row_data)  # 打印该行内容
            table_context_list.append(row_data)  # 将该行内容添加到表格内容列表中
            row_index += 1
            # 限制表格解析行数，防止上下文溢出
            if row_index > 4:
                break
        temp_list = []
        for row_data in table_context_list:
            row_text = ','.join(row_data)
            row_text = f'[{row_text}]'
            temp_list.append(row_text)
        table_context = '\n'.join(temp_list)
        # 返回对象结构
        table_dic = {'table_index':table_index, 'table':table,'table_context':table_context,'doc_meta_dic':doc_meta_dic}
        table_dic_list.append(table_dic)
    
    return table_dic_list


def merge_db_with_info_to_mongodb(table_objects,doc_meta_dic,table_index):
    '''保存表格数据到MongoDB'''
    for obj in table_objects:
        db_data_dic = transform_keys(obj, key_mapping)
        # 新增数据写入时间
        extend_info = {'created_time': time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())}
        # 合并文档元数据
        merge_dic = db_data_dic | doc_meta_dic | {'table_index': table_index} | extend_info
        try:                
            save_to_mongodb(merge_dic)
        except Exception as e:
            print(f"存储到MongoDB失败: {str(e)}")            

def retrieve_table_info_with_llm(table_index,table,table_context,doc_meta_dic):
    if is_debug:
        print(f"开始解析表格{table_index}...")
    llm_result = search_table_header_with_llm(table_context)
    if llm_result is None:
        print("LLM返回空结果，跳过！")
        log_to_mongodb({'level':'error','message':f"LLM解析表格{table_index}失败，文件路径：{doc_meta_dic['file_path']}"})
        return
    xml_result,start_row = extract_xml_to_dict(llm_result)
    if start_row is None:
        log_to_mongodb({'level':'error','message':f"XML中{start_row}不存在，解析失败，文件路径：{doc_meta_dic['file_path']}"})
        return
    if xml_result is not None:
        is_xml_valid = is_valid_data(xml_result)
    else:
        is_xml_valid = False
    if is_xml_valid:
        if is_debug:
            print(f"LLM返回结果：\n{llm_result}")
            #print(f"表格内容：\n{table_context}")
            print(f"XML结构：\n{xml_result}")
            #print(f"start_row：\n{start_row}")
        table_objects = parse_table_to_objects(table, xml_result, start_row)
        if is_debug:
            pass
            #print(f"对象列表：\n{table_objects}")
        merge_db_with_info_to_mongodb(table_objects,doc_meta_dic,table_index)
    elif xml_result is not None and is_xml_valid is False:
        # print(xml_result)
        # 需要进入重试机制，重新分析表格内容
        print("LLM解析非有效XML结构，重新分析！")
        # LLM重新数据分析
        llm_result = search_table_header_with_llm(table_context)
        xml_result,start_row = extract_xml_to_dict(llm_result)
        if xml_result is not None:
            is_xml_valid = is_valid_data(xml_result)
        else:
            is_xml_valid = False
            log_to_mongodb({'level':'info','message':f"解析表格{table_index}失败，文件路径：{doc_meta_dic['file_path']}"})
            print("LLM重新解析失败！")
        if is_xml_valid:
            # 重新解析后成功了再保存
            table_objects = parse_table_to_objects(table, xml_result, start_row)
            merge_db_with_info_to_mongodb(table_objects,doc_meta_dic,table_index)
    else:
        # print(f"表格内容：\n{table_context}")
        # print(f"llm_result:\n{llm_result}")
        # print("无XML结构!")
        return


def parallel(arg_list,func,max_workers=4,show_progress=False):
    '''并行处理任务，等待返回结果'''
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        """Process multiple inputs concurrently with the same prompt."""
        futures = [executor.submit(func,**item) for item in arg_list]
        if show_progress:
            # 使用 tqdm 包装 futures，实时显示进度
            results = [f.result() for f in tqdm(futures, total=len(futures), desc="Processing")]
        else:
            results = [f.result() for f in futures]
        return results

#TODO: 需要优化子任务也是异步的通信机制
def process_list_concurrently(items_list,process_func, max_workers=4):
    '''并行处理任务，不需要要等待返回结果'''
    start_time = time.time()
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # 提交所有任务，不等待结果
        executor.map(process_func, items_list)
    
    end_time = time.time()
    print(f"处理完成，耗时：{end_time-start_time}秒")

def parse_docx_tables(file_path):
    print(f'开始处理文档：{file_path}')
    if not custom_filter(file_path):
        # print(f"文件{file_path}不符合过滤条件，跳过！")
        log_to_mongodb({'level':'info','type':'file_filtered','message':f"文件{file_path}不符合过滤条件，跳过！"})
        return f'{file_path}不符合过滤条件，跳过！'
    table_dic_list = retrieve_table_from_docx(file_path)
    parallel(table_dic_list,retrieve_table_info_with_llm,max_workers=10)
    
    try:
        #处理完这个文件后添加到已处理文件中，并记录日志
        # destination_path = destination_path
        # 移动文件
        shutil.move(file_path, destination_path)
        log_to_mongodb({'level':'info','type':'file_resolved','message':f"文件解析处理完成，文件路径：{file_path}"})
        return f'文件{file_path}解析处理完成！'
    except Exception as e:
        # print(f"文件处理失败：{str(e)}")
        log_to_mongodb({'level':'error','type':'file_resolved','message':f"文件{file_path}处理失败，原因：{str(e)}"})
        return f'文件{file_path}处理完成！但是移动失败！'

def get_all_files(directory):
    '''递归获取所有文件（包括子文件夹）'''
    all_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            all_files.append(os.path.join(root, file))
    return all_files

def custom_filter(file_path:str):
    '''自定义文件过滤器'''
    if "技术协议" in file_path:
        return False    
    if file_path.endswith('.docx'):
        return True
    return False

def test_single_file():
    parse_docx_tables(file_path)

def test_batch_files(worker_num):
    all_files = get_all_files(source_path)
    all_files_dic_list = [{'file_path':file} for file in all_files]
    print(f"共{len(all_files)}个文件")
    parallel(all_files_dic_list,parse_docx_tables,max_workers=worker_num,show_progress=True)


def main():
    start_time = time.time()
    # test_single_file()
    test_batch_files(worker_num=10)
    end_time = time.time()
    print(f"总耗时：{end_time-start_time}秒")


if __name__ == '__main__':
    
    load_dotenv()
    # 获取环境变量
    deepseek_base_url = os.getenv("DEEPSEEK_BASE_URL")
    deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
    # deepseek_base_url = os.getenv("OPENROUTER_BASE_URL")
    # deepseek_api_key = os.getenv("OPENROUTER_API_KEY")

       
    is_debug = False
    # model_name='deepseek/deepseek-chat-v3-0324:free'
    model_name='deepseek-chat'
    # model_name='deepseek-reasoner'
    source_path = r'C:\Lee\work\contract\all\trans'
    destination_path = r'C:\Lee\work\contract\all\processed'
    # source_path = r'C:\Lee\files\采购\others'
    # destination_path = r'C:\Lee\files\采购\processed'
    main()
   